import numpy as np
from docx import Document

# ------------------------------
# Helpers
# ------------------------------
def embed_text(texts, client):
    response = client.embeddings.create(
        model="text-embedding-004",
        input=texts
    )
    return [np.array(r.embedding, dtype="float32") for r in response.data]

def chunk_text(text, chunk_size=500, overlap=50):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
    return chunks

def generate_answer(query, retrieved_chunks, chat_history, client):
    history_str = "\n".join(
        [f"{msg['role'].capitalize()}: {msg['content']}" for msg in chat_history if msg['role'] != "sources"]
    )
    context = "\n\n".join([f"From {c['doc_name']}:\n{c['text']}" for c in retrieved_chunks])

    prompt = f"""You are a helpful assistant specialized in healthcare documents.
Use only the provided context to answer.
If something is not in the context, say you don't know. Handle medical terminology and abbreviations correctly.

Conversation so far:
{history_str}

Query: {query}

Context:
{context}

Answer:"""

    response = client.chat.completions.create(
        model="gemini-1.5-flash",
        messages=[
            {"role": "system", "content": "You are a helpful AI assistant specialized in healthcare documents."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )

    return response.choices[0].message.content

def extract_relevant_text(chunk_text, query, client):
    """Use LLM to extract only relevant part of a chunk for the query."""
    prompt = f"""
You are an assistant specialized in healthcare documents. 
Extract only the sentences or phrases from the text below that are relevant to the question.

Question: {query}

Text: {chunk_text}

Return only the relevant part.
"""
    response = client.chat.completions.create(
        model="gemini-1.5-flash",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )
    return response.choices[0].message.content.strip()

def extract_text_from_docx(file_path):
    """Extract plain text from DOCX including paragraphs and tables."""
    doc = Document(file_path)
    texts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                texts.append(row_text)

    full_text = "\n".join(texts)
    return full_text
